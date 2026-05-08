from __future__ import annotations

import csv
import hashlib
import html as html_lib
import json
import re
import sys
import time
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Any
from urllib.parse import urldefrag, urljoin, urlparse
from urllib.robotparser import RobotFileParser
from urllib.error import HTTPError

try:
    from scrapling.fetchers import Fetcher
except Exception as exc:  # pragma: no cover - user environment dependency
    Fetcher = None
    SCRAPLING_IMPORT_ERROR = exc
else:
    SCRAPLING_IMPORT_ERROR = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except Exception as exc:  # pragma: no cover - user environment dependency
    Document = None
    DOCX_IMPORT_ERROR = exc
else:
    DOCX_IMPORT_ERROR = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
except Exception as exc:  # pragma: no cover - user environment dependency
    SimpleDocTemplate = None
    PDF_IMPORT_ERROR = exc
else:
    PDF_IMPORT_ERROR = None


ROOT = Path(__file__).resolve().parent
CONFIG_PATH = ROOT / "config.json"
DATA_DIR = ROOT / "data"
DAILY_DIR = DATA_DIR / "daily"
CACHE_DIR = ROOT / ".cache" / "html"
REPORTS_DIR = ROOT / "reports"
HISTORY_JSON = DATA_DIR / "drap_updates_history.json"
HISTORY_CSV = DATA_DIR / "drap_updates_history.csv"
SEEN_IDS_JSON = DATA_DIR / "seen_item_ids.json"
USER_AGENT = "DRAPDailyMonitor/1.0 (+public low-rate monitoring)"


@dataclass
class LinkItem:
    title: str
    url: str
    source_url: str
    matched_by: str
    classification: str
    reason: str
    content_type: str
    content_excerpt: str
    content_text: str
    first_seen: str
    last_seen: str
    item_id: str


@dataclass
class CrawledPage:
    title: str
    url: str
    source_url: str
    matched_by: str
    classification: str
    reason: str
    content_type: str
    content_excerpt: str
    content_text: str
    first_seen: str
    last_seen: str
    item_id: str


class LinkExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.links: list[dict[str, str]] = []
        self.page_title_parts: list[str] = []
        self._current_href: str | None = None
        self._current_text: list[str] = []
        self._in_title = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() == "a":
            attr_map = {key.lower(): value or "" for key, value in attrs}
            href = attr_map.get("href")
            if href:
                self._current_href = href
                self._current_text = []
        elif tag.lower() == "title":
            self._in_title = True

    def handle_data(self, data: str) -> None:
        if self._current_href:
            self._current_text.append(data)
        if self._in_title:
            self.page_title_parts.append(data)

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() == "a" and self._current_href:
            title = clean_text(" ".join(self._current_text))
            self.links.append({"href": self._current_href, "title": title})
            self._current_href = None
            self._current_text = []
        elif tag.lower() == "title":
            self._in_title = False


class TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []
        self._skip_depth = 0
        self._block_tag = False

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg", "form", "nav", "footer"}:
            self._skip_depth += 1
        if tag in {"p", "h1", "h2", "h3", "h4", "h5", "li", "td", "th", "article", "section", "br"}:
            self._block_tag = True

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in {"script", "style", "noscript", "svg", "form", "nav", "footer"} and self._skip_depth:
            self._skip_depth -= 1
        if tag in {"p", "h1", "h2", "h3", "h4", "h5", "li", "td", "th", "article", "section", "br"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        if self._skip_depth:
            return
        text = clean_text(data)
        if not text:
            return
        if self._block_tag and self.parts and self.parts[-1] != "\n":
            self.parts.append("\n")
        self.parts.append(text)
        self._block_tag = False


def clean_text(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def load_config() -> dict[str, Any]:
    with CONFIG_PATH.open("r", encoding="utf-8-sig") as handle:
        return json.load(handle)


def ensure_dirs() -> None:
    for path in (DATA_DIR, DAILY_DIR, CACHE_DIR, REPORTS_DIR):
        path.mkdir(parents=True, exist_ok=True)


def today_stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d")


def run_stamp() -> str:
    return datetime.now().strftime("%Y-%m-%d_%H%M%S")


def cache_path_for(url: str) -> Path:
    digest = hashlib.sha256(url.encode("utf-8")).hexdigest()[:16]
    return CACHE_DIR / f"{digest}.html"


def normalize_url(url: str) -> str:
    clean, _fragment = urldefrag(url)
    parsed = urlparse(clean)
    if parsed.scheme not in ("http", "https"):
        return clean
    if parsed.path != "/" and clean.endswith("/"):
        clean = clean.rstrip("/")
    return clean


def domain_allowed(url: str, allowed_domains: list[str]) -> bool:
    host = (urlparse(url).hostname or "").lower()
    return host in {domain.lower() for domain in allowed_domains}


def robots_allowed(url: str, enabled: bool) -> tuple[bool, str]:
    if not enabled:
        return True, "robots.txt disabled in config"

    parsed = urlparse(url)
    robots_url = f"{parsed.scheme}://{parsed.netloc}/robots.txt"
    parser = RobotFileParser()
    parser.set_url(robots_url)
    try:
        parser.read()
    except Exception as exc:
        return True, f"robots.txt could not be read, proceeding cautiously: {exc}"

    allowed = parser.can_fetch(USER_AGENT, url)
    reason = "allowed by robots.txt" if allowed else "disallowed by robots.txt"
    return allowed, reason


def page_to_html(page: Any) -> str:
    for attr in ("html", "text", "body", "content"):
        if hasattr(page, attr):
            value = getattr(page, attr)
            value = value() if callable(value) else value
            if isinstance(value, bytes):
                return value.decode("utf-8", errors="replace")
            if isinstance(value, str):
                return value
    return str(page)


def fetch_html(url: str, config: dict[str, Any]) -> tuple[str, str]:
    cached = cache_path_for(url)
    if config.get("development_cache") and cached.exists() and cached.stat().st_size > 100:
        return cached.read_text(encoding="utf-8", errors="replace"), "cache"

    html = ""
    source = "network"

    if Fetcher is not None:
        page = Fetcher.get(url)
        html = page_to_html(page)
        source = "scrapling"

    if len(clean_text(html)) < 100:
        request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
        with urllib.request.urlopen(request, timeout=30) as response:
            html = response.read().decode("utf-8", errors="replace")
        source = "urllib-fallback" if Fetcher is not None else "urllib"

    if len(clean_text(html)) < 100:
        raise RuntimeError(f"Fetched empty or near-empty response from {url}")

    if config.get("development_cache"):
        cached.write_text(html, encoding="utf-8")
    return html, source


def detect_blocked(html: str, config: dict[str, Any]) -> tuple[bool, str]:
    blocked_config = config.get("blocked_detection", {})
    if not blocked_config.get("enabled", True):
        return False, ""
    lowered = html.lower()
    for marker in blocked_config.get("markers", []):
        if marker.lower() in lowered:
            return True, marker
    return False, ""


def extract_links(html: str, source_url: str) -> list[dict[str, str]]:
    parser = LinkExtractor()
    parser.feed(html)
    results = []
    for raw in parser.links:
        href = raw["href"]
        title = raw["title"] or href
        if href.startswith("#") or href.lower().startswith(("javascript:", "mailto:", "tel:")):
            continue
        results.append({"title": title, "url": normalize_url(urljoin(source_url, href))})
    return results


def extract_page_title(html: str, fallback_url: str) -> str:
    parser = LinkExtractor()
    parser.feed(html)
    title = clean_text(" ".join(parser.page_title_parts))
    if title:
        return title
    path = urlparse(fallback_url).path.strip("/")
    return path or fallback_url


def extract_page_text(html: str, limit: int | None = None) -> str:
    parser = TextExtractor()
    parser.feed(html)
    text = "\n".join(
        line for line in (clean_text(part) for part in "".join(parser.parts).splitlines())
        if line
    )
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if limit and len(text) > limit:
        return text[:limit].rsplit(" ", 1)[0] + "..."
    return text


def classify(title: str, url: str, config: dict[str, Any]) -> tuple[str, str]:
    text = f"{title} {url}".lower()
    classification = config.get("classification", {})

    for word in classification.get("bad", []):
        if word.lower() in text:
            return "bad", f"matched bad keyword: {word}"

    for word in classification.get("good", []):
        if word.lower() in text:
            return "good", f"matched good keyword: {word}"

    return "neutral", "no good/bad keyword matched"


def flexible_match(title: str, url: str, config: dict[str, Any]) -> tuple[bool, str]:
    if config.get("record_all_internal_links", False):
        return True, "all internal links mode"

    selector = config.get("selector_strategy", {})
    title_lower = title.lower()
    url_lower = url.lower()

    for word in selector.get("title_keywords", []):
        if word.lower() in title_lower:
            return True, f"title keyword: {word}"

    for word in selector.get("url_keywords", []):
        if word.lower() in url_lower:
            return True, f"url keyword: {word}"

    for pattern in selector.get("regex_patterns", []):
        if re.search(pattern, f"{title} {url}"):
            return True, f"regex: {pattern}"

    if selector.get("smart_element_tracking") and is_probably_update_link(title, url):
        return True, "smart similarity fallback"

    return False, ""


def is_probably_update_link(title: str, url: str) -> bool:
    text = f"{title} {url}".lower()
    updateish = ("pdf" in text or "download" in text or "uploads" in text)
    has_date = bool(re.search(r"\b20\d{2}\b|\b\d{1,2}[-_/]\d{1,2}[-_/]\d{2,4}\b", text))
    return updateish and has_date


def make_item_id(title: str, url: str) -> str:
    return hashlib.sha256(f"{title}|{url}".encode("utf-8")).hexdigest()[:20]


def path_extension(url: str) -> str:
    path = urlparse(url).path.lower()
    if "." not in path.rsplit("/", 1)[-1]:
        return ""
    return "." + path.rsplit(".", 1)[-1]


def should_crawl_url(url: str, config: dict[str, Any]) -> bool:
    ext = path_extension(url)
    if ext in {item.lower() for item in config.get("record_file_extensions", [])}:
        return False
    if ext and ext not in {item.lower() for item in config.get("crawl_file_extensions", [])}:
        return False
    return True


def should_record_url(url: str, config: dict[str, Any]) -> bool:
    ext = path_extension(url)
    if ext in {item.lower() for item in config.get("record_file_extensions", [])}:
        return True
    return True


def load_history() -> list[dict[str, Any]]:
    if not HISTORY_JSON.exists():
        return []
    with HISTORY_JSON.open("r", encoding="utf-8") as handle:
        return json.load(handle)


def load_seen_ids() -> set[str]:
    if SEEN_IDS_JSON.exists():
        with SEEN_IDS_JSON.open("r", encoding="utf-8") as handle:
            return set(json.load(handle))
    return {row["item_id"] for row in load_history() if row.get("item_id")}


def save_seen_ids(item_ids: set[str]) -> None:
    save_json(SEEN_IDS_JSON, sorted(item_ids))


def save_json(path: Path, rows: list[dict[str, Any]]) -> None:
    path.write_text(json.dumps(rows, indent=2, ensure_ascii=False), encoding="utf-8")


def save_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    fields = [
        "first_seen",
        "last_seen",
        "classification",
        "title",
        "url",
        "source_url",
        "matched_by",
        "reason",
        "content_type",
        "content_excerpt",
        "item_id",
    ]
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def api_get_json(url: str) -> tuple[Any, dict[str, str]]:
    request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT})
    with urllib.request.urlopen(request, timeout=60) as response:
        data = json.loads(response.read().decode("utf-8", errors="replace"))
        headers = {key: value for key, value in response.headers.items()}
    return data, headers


def unescape_text(value: str) -> str:
    return clean_text(html_lib.unescape(value or ""))


def collect_wordpress_api(config: dict[str, Any]) -> tuple[list[LinkItem | CrawledPage], dict[str, Any]]:
    api_config = config.get("wordpress_api", {})
    per_page = int(api_config.get("per_page", 100))
    content_types = api_config.get("content_types", ["posts", "pages"])
    found: list[LinkItem | CrawledPage] = []
    api_items = 0
    failed_pages = 0
    links_discovered = 0
    blocked_reason = "none"

    base = config["start_urls"][0].rstrip("/")
    endpoint_map = {
        "posts": "posts",
        "pages": "pages",
    }

    for content_type in content_types:
        endpoint = endpoint_map.get(content_type, content_type)
        page_number = 1
        total_pages = None

        while total_pages is None or page_number <= total_pages:
            api_url = f"{base}/wp-json/wp/v2/{endpoint}?per_page={per_page}&page={page_number}"
            try:
                records, headers = api_get_json(api_url)
            except HTTPError as exc:
                if exc.code == 400 and page_number > 1:
                    break
                failed_pages += 1
                blocked_reason = f"API page failed: {api_url} ({exc})"
                break
            except Exception as exc:
                failed_pages += 1
                blocked_reason = f"API page failed: {api_url} ({exc})"
                break

            total_pages = int(headers.get("X-WP-TotalPages", "1"))
            api_items += len(records)

            for record in records:
                title = unescape_text(record.get("title", {}).get("rendered", "")) or f"{content_type} {record.get('id')}"
                url = normalize_url(record.get("link", ""))
                content_html = record.get("content", {}).get("rendered", "") or record.get("excerpt", {}).get("rendered", "")
                full_text = extract_page_text(content_html)
                excerpt = extract_page_text(content_html, int(config.get("report_text_chars", 1000)))
                classification, reason = classify(f"{title} {full_text[:500]}", url, config)

                found.append(CrawledPage(
                    title=title,
                    url=url,
                    source_url=api_url,
                    matched_by=f"wordpress REST API {content_type}",
                    classification=classification,
                    reason=reason,
                    content_type=f"wordpress-{content_type.rstrip('s')}",
                    content_excerpt=excerpt,
                    content_text=full_text,
                    first_seen=today_stamp(),
                    last_seen=today_stamp(),
                    item_id=make_item_id(title, url),
                ))

                for link in extract_links(content_html, url):
                    if not domain_allowed(link["url"], config.get("allowed_domains", [])):
                        continue
                    links_discovered += 1
                    ext = path_extension(link["url"])
                    if ext not in {item.lower() for item in config.get("record_file_extensions", [])}:
                        continue
                    link_classification, link_reason = classify(link["title"], link["url"], config)
                    found.append(LinkItem(
                        title=link["title"] or link["url"],
                        url=link["url"],
                        source_url=url,
                        matched_by=f"linked upload in {content_type}",
                        classification=link_classification,
                        reason=link_reason,
                        content_type="linked-file",
                        content_excerpt="",
                        content_text="",
                        first_seen=today_stamp(),
                        last_seen=today_stamp(),
                        item_id=make_item_id(link["title"], link["url"]),
                    ))

            print(f"Fetched {content_type} API page {page_number}/{total_pages} ({api_items} items so far)")
            page_number += 1

    run_status = {
        "status": "successful" if failed_pages == 0 else "completed with skipped errors",
        "pages_checked": api_items,
        "failed_pages": failed_pages,
        "links_discovered": links_discovered,
        "blocked_reason": blocked_reason,
    }
    return found, run_status


def merge_history(existing: list[dict[str, Any]], found: list[LinkItem | CrawledPage]) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    today = today_stamp()
    by_id = {item["item_id"]: item for item in existing}
    new_rows = []

    for item in found:
        row = item.__dict__.copy()
        if item.item_id in by_id:
            by_id[item.item_id]["last_seen"] = today
        else:
            by_id[item.item_id] = row
            new_rows.append(row)

    history = sorted(by_id.values(), key=lambda row: (row["first_seen"], row["title"]), reverse=True)
    return history, new_rows


def split_new_items(found: list[LinkItem | CrawledPage], seen_ids: set[str]) -> list[dict[str, Any]]:
    new_rows = []
    for item in found:
        if item.item_id in seen_ids:
            continue
        new_rows.append(item.__dict__.copy())
    return new_rows


def dedupe_items(items: list[LinkItem | CrawledPage]) -> list[LinkItem | CrawledPage]:
    seen = set()
    unique = []
    for item in items:
        if item.item_id in seen:
            continue
        seen.add(item.item_id)
        unique.append(item)
    return unique


def pdf_text(value: Any) -> str:
    return html_lib.escape(str(value or ""))


def build_docx_report(report_path: Path, new_rows: list[dict[str, Any]], run_status: dict[str, Any]) -> None:
    if Document is None:
        raise RuntimeError(f"python-docx could not import: {DOCX_IMPORT_ERROR}")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DRAP Daily Update Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(32, 74, 112)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Date checked: {today_stamp()}").italic = True

    doc.add_heading("Run Summary", level=1)
    summary = doc.add_table(rows=1, cols=2)
    summary.style = "Table Grid"
    summary.rows[0].cells[0].text = "Status"
    summary.rows[0].cells[1].text = run_status["status"]
    for key in ("pages_checked", "failed_pages", "links_discovered", "items_in_report", "new_updates", "blocked_reason"):
        cells = summary.add_row().cells
        cells[0].text = key.replace("_", " ").title()
        cells[1].text = str(run_status.get(key, ""))

    grouped = {
        "bad": [row for row in new_rows if row["classification"] == "bad"],
        "good": [row for row in new_rows if row["classification"] == "good"],
        "neutral": [row for row in new_rows if row["classification"] == "neutral"],
    }

    for section_name in ("bad", "good", "neutral"):
        doc.add_heading(f"{section_name.title()} Items", level=1)
        rows = grouped[section_name]
        if not rows:
            doc.add_paragraph("No new items in this category.")
            continue
        for index, row in enumerate(rows, start=1):
            para = doc.add_paragraph()
            para.add_run(f"{index}. {row['title']}").bold = True
            doc.add_paragraph(f"Type: {row.get('content_type', 'unknown')}")
            doc.add_paragraph(f"Reason: {row['reason']}")
            doc.add_paragraph(f"Source: {row['url']}")
            excerpt = row.get("content_excerpt") or ""
            if excerpt:
                doc.add_paragraph("Content:")
                doc.add_paragraph(excerpt)

    doc.save(report_path)


def build_pdf_report(report_path: Path, new_rows: list[dict[str, Any]], run_status: dict[str, Any]) -> None:
    if SimpleDocTemplate is None:
        raise RuntimeError(f"reportlab could not import: {PDF_IMPORT_ERROR}")

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10))
    story: list[Any] = []

    story.append(Paragraph("DRAP Daily Update Report", styles["Title"]))
    story.append(Paragraph(f"Date checked: {today_stamp()}", styles["BodyText"]))
    story.append(Spacer(1, 0.2 * inch))

    summary_rows = [
        ["Status", run_status["status"]],
        ["Pages checked", str(run_status.get("pages_checked", ""))],
        ["Failed/skipped pages", str(run_status.get("failed_pages", ""))],
        ["Links discovered", str(run_status.get("links_discovered", ""))],
        ["Items in report", str(run_status.get("items_in_report", ""))],
        ["New updates", str(run_status.get("new_updates", ""))],
        ["Blocked reason", str(run_status.get("blocked_reason", ""))],
    ]
    table = Table(summary_rows, colWidths=[1.6 * inch, 4.8 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#E8F0F7")),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C7D4")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.25 * inch))

    for section_name in ("bad", "good", "neutral"):
        story.append(Paragraph(f"{section_name.title()} Items", styles["Heading2"]))
        rows = [row for row in new_rows if row["classification"] == section_name]
        if not rows:
            story.append(Paragraph("No new items in this category.", styles["BodyText"]))
            continue
        for index, row in enumerate(rows, start=1):
            story.append(Paragraph(f"{index}. {pdf_text(row['title'])}", styles["Heading3"]))
            story.append(Paragraph(f"Type: {pdf_text(row.get('content_type', 'unknown'))}", styles["BodyText"]))
            story.append(Paragraph(f"Reason: {pdf_text(row['reason'])}", styles["BodyText"]))
            story.append(Paragraph(f"Source: {pdf_text(row['url'])}", styles["Small"]))
            excerpt = row.get("content_excerpt") or ""
            if excerpt:
                story.append(Paragraph("Content:", styles["BodyText"]))
                story.append(Paragraph(pdf_text(excerpt).replace("\n", "<br/>"), styles["Small"]))
            story.append(Spacer(1, 0.1 * inch))

    doc = SimpleDocTemplate(str(report_path), pagesize=A4, rightMargin=42, leftMargin=42, topMargin=42, bottomMargin=42)
    doc.build(story)


def run() -> int:
    config = load_config()
    ensure_dirs()
    if config.get("wordpress_api", {}).get("enabled", False):
        found, api_status = collect_wordpress_api(config)
        found = dedupe_items(found)
        seen_ids = load_seen_ids()
        new_rows = split_new_items(found, seen_ids)
        report_rows = new_rows if config.get("report_scope") == "new_only" else [item.__dict__.copy() for item in found]
        seen_ids.update(item.item_id for item in found)
        save_seen_ids(seen_ids)
        run_status = {
            **api_status,
            "items_in_report": len(report_rows),
            "new_updates": len(new_rows),
            "max_pages_per_run": "wordpress_api",
        }

        daily_base = DAILY_DIR / f"{today_stamp()}_new_updates"
        save_json(daily_base.with_suffix(".json"), report_rows)
        save_csv(daily_base.with_suffix(".csv"), report_rows)

        report_base = REPORTS_DIR / f"{run_stamp()}_DRAP_new_updates"
        build_docx_report(report_base.with_suffix(".docx"), report_rows, run_status)
        build_pdf_report(report_base.with_suffix(".pdf"), report_rows, run_status)

        print(json.dumps(run_status, indent=2))
        print(f"Report DOCX: {report_base.with_suffix('.docx')}")
        print(f"Report PDF:  {report_base.with_suffix('.pdf')}")
        return 0

    checked_pages = 0
    links_discovered = 0
    failed_pages = 0
    found: list[LinkItem | CrawledPage] = []
    blocked_reason = ""
    status = "successful"
    max_pages = int(config.get("max_pages_per_run", 300))
    queue = [normalize_url(url) for url in config["start_urls"]]
    queued = set(queue)
    visited: set[str] = set()

    while queue and checked_pages < max_pages:
        url = queue.pop(0)
        if url in visited:
            continue
        visited.add(url)

        if not domain_allowed(url, config.get("allowed_domains", [])):
            continue
        if not should_crawl_url(url, config):
            continue

        allowed, robot_reason = robots_allowed(url, config.get("robots_txt_obey", True))
        if not allowed:
            status = "skipped by robots.txt"
            blocked_reason = robot_reason
            continue

        try:
            html, source = fetch_html(url, config)
        except Exception as exc:
            failed_pages += 1
            blocked_reason = f"{failed_pages} page(s) failed, latest: {url} ({exc})"
            if status == "successful":
                status = "completed with skipped errors"
            continue
        checked_pages += 1

        blocked, marker = detect_blocked(html, config)
        if blocked:
            status = "manual check needed"
            blocked_reason = f"blocked/CAPTCHA marker detected: {marker}"
            break

        page_title = extract_page_title(html, url)
        page_text = extract_page_text(html)
        page_excerpt = extract_page_text(html, int(config.get("report_text_chars", 2500)))
        page_classification, page_reason = classify(page_title, url, config)
        found.append(CrawledPage(
            title=page_title,
            url=url,
            source_url=url,
            matched_by=f"crawled page; fetched from {source}",
            classification=page_classification,
            reason=page_reason,
            content_type="html-page",
            content_excerpt=page_excerpt,
            content_text=page_text,
            first_seen=today_stamp(),
            last_seen=today_stamp(),
            item_id=make_item_id(page_title, url),
        ))

        links = extract_links(html, url)
        links_discovered += len(links)
        for link in links:
            if not domain_allowed(link["url"], config.get("allowed_domains", [])):
                continue
            if should_crawl_url(link["url"]) and link["url"] not in queued and link["url"] not in visited:
                queue.append(link["url"])
                queued.add(link["url"])

            if not should_record_url(link["url"], config):
                continue
            matched, matched_by = flexible_match(link["title"], link["url"], config)
            if not matched:
                continue
            classification, reason = classify(link["title"], link["url"], config)
            found.append(LinkItem(
                title=link["title"],
                url=link["url"],
                source_url=url,
                matched_by=f"{matched_by}; fetched from {source}",
                classification=classification,
                reason=reason,
                content_type="linked-file" if path_extension(link["url"]) else "internal-link",
                content_excerpt="",
                content_text="",
                first_seen=today_stamp(),
                last_seen=today_stamp(),
                item_id=make_item_id(link["title"], link["url"]),
            ))

        time.sleep(float(config.get("request_delay_seconds", 3)))

    found = dedupe_items(found)
    seen_ids = load_seen_ids()
    new_rows = split_new_items(found, seen_ids)
    report_rows = new_rows if config.get("report_scope") == "new_only" else [item.__dict__.copy() for item in found]
    seen_ids.update(item.item_id for item in found)
    save_seen_ids(seen_ids)
    run_status = {
        "status": status,
        "pages_checked": checked_pages,
        "failed_pages": failed_pages,
        "links_discovered": links_discovered,
        "items_in_report": len(report_rows),
        "new_updates": len(new_rows),
        "max_pages_per_run": max_pages,
        "blocked_reason": blocked_reason or "none",
    }

    daily_base = DAILY_DIR / f"{today_stamp()}_new_updates"
    save_json(daily_base.with_suffix(".json"), report_rows)
    save_csv(daily_base.with_suffix(".csv"), report_rows)

    report_base = REPORTS_DIR / f"{run_stamp()}_DRAP_new_updates"
    build_docx_report(report_base.with_suffix(".docx"), report_rows, run_status)
    build_pdf_report(report_base.with_suffix(".pdf"), report_rows, run_status)

    print(json.dumps(run_status, indent=2))
    print(f"Report DOCX: {report_base.with_suffix('.docx')}")
    print(f"Report PDF:  {report_base.with_suffix('.pdf')}")
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(run())
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
